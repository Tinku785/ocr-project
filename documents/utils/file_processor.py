import os
import csv
import tempfile
from docx import Document as DocxDocument
from openpyxl import load_workbook


def get_file_type(filename):
    ext = os.path.splitext(filename)[1].lower()
    if ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
        return 'image'
    elif ext == '.pdf':
        return 'pdf'
    elif ext == '.docx':
        return 'docx'
    elif ext == '.xlsx':
        return 'xlsx'
    elif ext == '.csv':
        return 'csv'
    return 'unknown'


def extract_from_pdf(file_path, doc_id=None):
    import fitz
    from .ocr_engine import extract_text
    from .validator import validate_text

    text = ''

    with fitz.open(file_path) as pdf:
        total_pages = len(pdf)
        for page_num in range(total_pages):
            if doc_id:
                try:
                    from ..models import Document
                    progress_pct = 10 + int((page_num / total_pages) * 80)
                    Document.objects.filter(pk=doc_id).update(progress=progress_pct)
                except Exception as e:
                    print(f"Error updating progress: {e}")
            page = pdf[page_num]

            # Try direct text extraction first
            page_text = page.get_text()

            if page_text.strip():
                # Digital PDF — text found directly
                print(f"Page {page_num + 1}: digital text found")
                text += page_text + '\n'
            else:
                # Scanned PDF — convert page to image and run OCR
                print(f"Page {page_num + 1}: no text found, running OCR")
                mat = fitz.Matrix(2, 2)  # 2x zoom for better quality
                pix = page.get_pixmap(matrix=mat)

                with tempfile.NamedTemporaryFile(
                    suffix='.png', delete=False
                ) as tmp:
                    pix.save(tmp.name)
                    tmp_path = tmp.name

                try:
                    raw_items = extract_text(tmp_path)
                    page_text = validate_text(raw_items)
                    if page_text.strip():
                        text += page_text + '\n'
                except Exception as e:
                    print(f"OCR failed for page {page_num + 1}: {e}")
                finally:
                    os.unlink(tmp_path)

    return text.strip()


def extract_from_docx(file_path):
    doc = DocxDocument(file_path)
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text.strip())
    # Also extract tables inside the docx
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                lines.append(row_text)
    return '\n'.join(lines)


def extract_from_xlsx(file_path):
    wb = load_workbook(file_path, data_only=True)
    lines = []
    for sheet in wb.worksheets:
        lines.append(f'Sheet: {sheet.title}')
        for row in sheet.iter_rows(values_only=True):
            row_text = ' | '.join(
                str(cell) for cell in row if cell is not None
            )
            if row_text.strip():
                lines.append(row_text)
    return '\n'.join(lines)


def extract_from_csv(file_path):
    lines = []
    with open(file_path, newline='', encoding='utf-8', errors='ignore') as f:
        reader = csv.reader(f)
        for row in reader:
            row_text = ' | '.join(cell.strip() for cell in row if cell.strip())
            if row_text:
                lines.append(row_text)
    return '\n'.join(lines)


def process_file(file_path, file_type, doc_id=None):
    if file_type == 'pdf':
        return extract_from_pdf(file_path, doc_id)
    elif file_type == 'docx':
        return extract_from_docx(file_path)
    elif file_type == 'xlsx':
        return extract_from_xlsx(file_path)
    elif file_type == 'csv':
        return extract_from_csv(file_path)
    return ''